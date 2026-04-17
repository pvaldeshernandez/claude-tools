#!/usr/bin/env python3
"""
Scientific Writer — Document Generator
Produces formatted .docx from structured input + journal profile + author database.

Supports two modes:
1. template_path in JSON: clone styles from a reference .docx (paragraphs are deleted,
   styles are preserved). This gives exact formatting from a real manuscript.
2. No template: build from scratch using journal profile YAML settings.

Usage:
    python3.13 docx-generator.py input.json

Input JSON schema:
{
    "title": "The paper title",
    "authors": ["montesino-goicolea", "valdes-hernandez", ...],
    "sections": [
        {"header": "Introduction", "text": "Body text...", "level": 1},
        {"header": "Participants", "text": "Body text...", "level": 2},
        ...
    ],
    "journal_profile": "journal-of-pain",
    "template_path": "/path/to/reference.docx",  // optional
    "output_path": "/path/to/output.docx"
}

Section "level" controls heading style:
  1 = Heading 1 (bold, major sections)
  2 = Heading 2 (bold italic, subsections)
  3 = Heading 3 (italic, sub-subsections)
  0 or omitted = no heading, just body text (e.g., Abstract body, Perspective)

Section "type" controls rendering:
  "text" (default) = normal body paragraphs
  "table_caption" = Subtitle style, "Table N." bold prefix
  "table" = tab-separated data → real Word table (Plain Table 1, 11pt, centered, autofit)
  "table_note" = Caption style (10pt), auto-prefixed with bold "Note: "
  "figure" = insert image inline + caption below (Caption style, "Figure N." bold)
  "figure_placeholder" = centered placeholder text (submission version)
  "figure_legends" = figure legend text at end of document

Optional top-level keys:
  "figures_dir": directory containing FigureN.png files (default: output dir)
"""
import json
import re
import sys
import os
import yaml
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
AUTHORS_FILE = os.path.join(SKILL_DIR, 'authors.yaml')
JOURNALS_DIR = os.path.join(SKILL_DIR, 'journals')
DEFAULT_TEMPLATE = os.path.join(SKILL_DIR, 'template.docx')

ALIGN_MAP = {
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
}

# ── Greek letter and math symbol mapping ─────────────────────────────
_GREEK_MAP = {
    'alpha': '\u03b1', 'beta': '\u03b2', 'gamma': '\u03b3', 'delta': '\u03b4',
    'epsilon': '\u03b5', 'zeta': '\u03b6', 'eta': '\u03b7', 'theta': '\u03b8',
    'iota': '\u03b9', 'kappa': '\u03ba', 'lambda': '\u03bb', 'mu': '\u03bc',
    'nu': '\u03bd', 'xi': '\u03be', 'pi': '\u03c0', 'rho': '\u03c1',
    'sigma': '\u03c3', 'tau': '\u03c4', 'upsilon': '\u03c5', 'phi': '\u03c6',
    'chi': '\u03c7', 'psi': '\u03c8', 'omega': '\u03c9',
    'varepsilon': '\u03b5', 'varphi': '\u03c6', 'vartheta': '\u03d1',
    'Alpha': '\u0391', 'Beta': '\u0392', 'Gamma': '\u0393', 'Delta': '\u0394',
    'Epsilon': '\u0395', 'Zeta': '\u0396', 'Eta': '\u0397', 'Theta': '\u0398',
    'Iota': '\u0399', 'Kappa': '\u039a', 'Lambda': '\u039b', 'Mu': '\u039c',
    'Nu': '\u039d', 'Xi': '\u039e', 'Pi': '\u03a0', 'Rho': '\u03a1',
    'Sigma': '\u03a3', 'Tau': '\u03a4', 'Upsilon': '\u03a5', 'Phi': '\u03a6',
    'Chi': '\u03a7', 'Psi': '\u03a8', 'Omega': '\u03a9',
}

_MATH_SYMBOLS = {
    'times': '\u00d7', 'cdot': '\u00b7', 'pm': '\u00b1', 'mp': '\u2213',
    'leq': '\u2264', 'geq': '\u2265', 'neq': '\u2260', 'approx': '\u2248',
    'sim': '\u223c', 'propto': '\u221d', 'infty': '\u221e', 'partial': '\u2202',
    'nabla': '\u2207', 'sum': '\u2211', 'prod': '\u220f', 'int': '\u222b',
    'sqrt': '\u221a', 'leftarrow': '\u2190', 'rightarrow': '\u2192',
    'to': '\u2192', 'mid': '|',
    'leftrightarrow': '\u2194', 'Leftarrow': '\u21d0', 'Rightarrow': '\u21d2',
    'in': '\u2208', 'notin': '\u2209', 'subset': '\u2282', 'supset': '\u2283',
    'cup': '\u222a', 'cap': '\u2229', 'forall': '\u2200', 'exists': '\u2203',
    'emptyset': '\u2205', 'neg': '\u00ac', 'wedge': '\u2227', 'vee': '\u2228',
}

# LaTeX operator names that should render as upright text in math
_MATH_OPERATORS = {
    'min', 'max', 'log', 'ln', 'exp', 'sin', 'cos', 'tan',
    'lim', 'sup', 'inf', 'det', 'dim', 'ker', 'mod', 'gcd',
}

# OMML accent characters for \bar, \hat, \tilde, \dot, \vec, etc.
_ACCENT_MAP = {
    'bar': '\u0305',     # combining overline
    'hat': '\u0302',     # combining circumflex
    'tilde': '\u0303',   # combining tilde
    'dot': '\u0307',     # combining dot above
    'ddot': '\u0308',    # combining diaeresis
    'vec': '\u20d7',     # combining right arrow above
    'check': '\u030c',   # combining caron
    'breve': '\u0306',   # combining breve
}

# Unicode superscript/subscript digit maps
_SUP_DIGITS = str.maketrans('0123456789+-=()', '\u2070\u00b9\u00b2\u00b3\u2074\u2075\u2076\u2077\u2078\u2079\u207a\u207b\u207c\u207d\u207e')
_SUB_DIGITS = str.maketrans('0123456789+-=()', '\u2080\u2081\u2082\u2083\u2084\u2085\u2086\u2087\u2088\u2089\u208a\u208b\u208c\u208d\u208e')


def load_authors():
    with open(AUTHORS_FILE, encoding='utf-8') as f:
        return yaml.safe_load(f)['collaborators']


def load_journal(profile_name):
    path = os.path.join(JOURNALS_DIR, f'{profile_name}.yaml')
    with open(path, encoding='utf-8') as f:
        return yaml.safe_load(f)


def compute_affiliations(author_keys, author_db):
    """Compute affiliation numbering for a specific set of authors."""
    seen = {}
    counter = 0
    authors_with_numbers = []

    for key in author_keys:
        info = author_db[key]
        nums = []
        for aff in info['affiliations']:
            if aff not in seen:
                counter += 1
                seen[aff] = counter
            nums.append(str(seen[aff]))
        authors_with_numbers.append((info['display'], ','.join(nums)))

    numbered_affiliations = [f'{num} {aff}' for aff, num in
                             sorted(seen.items(), key=lambda x: x[1])]
    return authors_with_numbers, numbered_affiliations


def _get_template_font(doc):
    """Extract the font name and size from the template's Normal style."""
    style = doc.styles['Normal']
    font_name = style.font.name or 'Times New Roman'
    font_size = style.font.size  # in EMU; None means inherited
    if font_size:
        font_size_pt = round(font_size / 12700)  # EMU to pt
    else:
        font_size_pt = 12
    return font_name, font_size_pt


def create_doc_from_template(template_path):
    """Open a template .docx, delete all content but keep styles."""
    doc = Document(template_path)
    body = doc.element.body
    # Remove all child elements from body except sectPr (page layout)
    sect_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'
    for child in list(body):
        if child.tag != sect_tag:
            body.remove(child)
    # Clear headers and footers
    for section in doc.sections:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        for p in section.header.paragraphs:
            p.text = ''
        for p in section.footer.paragraphs:
            p.text = ''
    # Clear footnotes and endnotes from the package
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for part_name in ['footnotes', 'endnotes']:
        try:
            part = doc.part.package.part_related_by(
                f'http://schemas.openxmlformats.org/officeDocument/2006/relationships/{part_name}')
            root = etree.fromstring(part.blob)
            # Keep only the required separator entries (id 0 and 1)
            for child in list(root):
                id_val = child.get(f'{{{ns}}}id', '')
                if id_val not in ('0', '1', '-1'):
                    root.remove(child)
            part._blob = etree.tostring(root, xml_declaration=True,
                                         encoding='UTF-8', standalone=True)
        except Exception:
            pass
    return doc


def create_doc_from_profile(journal):
    """Create a new doc and set styles from journal profile."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = journal.get('font', 'Times New Roman')
    style.font.size = Pt(journal.get('font_size', 12))
    alignment = journal.get('alignment', 'justify')
    style.paragraph_format.alignment = ALIGN_MAP.get(alignment,
                                                      WD_ALIGN_PARAGRAPH.JUSTIFY)
    if journal.get('space_before'):
        style.paragraph_format.space_before = Pt(journal['space_before'])
    if journal.get('space_after'):
        style.paragraph_format.space_after = Pt(journal['space_after'])
    if journal.get('line_spacing') == 'double':
        style.paragraph_format.line_spacing = 2.0
    if journal.get('first_line_indent'):
        style.paragraph_format.first_line_indent = Inches(
            journal['first_line_indent'])
    if journal.get('margins'):
        margins = journal['margins']
        for sec in doc.sections:
            if isinstance(margins, dict):
                def _parse_margin(val):
                    s = str(val)
                    if s.endswith('cm'):
                        return Cm(float(s.replace('cm', '')))
                    elif s.endswith('in'):
                        return Inches(float(s.replace('in', '')))
                    else:
                        return Inches(float(s))
                sec.left_margin = _parse_margin(margins.get('left', '1in'))
                sec.right_margin = _parse_margin(margins.get('right', '1in'))
                sec.top_margin = _parse_margin(margins.get('top', '1in'))
                sec.bottom_margin = _parse_margin(margins.get('bottom', '1in'))
            else:
                margin = Inches(float(str(margins).replace('in', '')))
                sec.top_margin = margin
                sec.bottom_margin = margin
                sec.left_margin = margin
                sec.right_margin = margin
    return doc


def add_heading(doc, text, level=1):
    """Add a heading using the document's heading styles."""
    return doc.add_heading(text, level=level)


_CITE_RE = re.compile(r'\[(\d+(?:,\d+)*(?:[-–]\d+)?)\]')

# Markdown-style formatting: **bold**, *italic*
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')
_ITALIC_RE = re.compile(r'\*(.+?)\*')
# Combined pattern that matches any inline formatting token
_FMT_RE = re.compile(
    r'(\[(?:\d+(?:,\d+)*(?:[-–]\d+)?)\])'   # citations [1,3]
    r'|(\*\*\*(?:.+?)\*\*\*)'                  # bold+italic ***text***
    r'|(\*\*(?:.+?)\*\*)'                      # bold **text**
    r'|(\*(?:.+?)\*)'                           # italic *text*
    r'|([A-Z][A-Za-z]*_[a-z]{1,5})'             # subscript: CWP_adj, F_max etc.
    r'|(https?://[^\s,;)\]]+|mailto:[^\s,;)\]]+)'  # URLs
)


def _add_hyperlink(p, url, display_text=None, font_name=None, font_size=None):
    """Add a clickable hyperlink to a paragraph (blue, underlined)."""
    part = p.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = etree.SubElement(p._element, qn('w:hyperlink'))
    hyperlink.set(qn('r:id'), r_id)
    run_elem = etree.SubElement(hyperlink, qn('w:r'))
    rPr = etree.SubElement(run_elem, qn('w:rPr'))
    rStyle = etree.SubElement(rPr, qn('w:rStyle'))
    rStyle.set(qn('w:val'), 'Hyperlink')
    color = etree.SubElement(rPr, qn('w:color'))
    color.set(qn('w:val'), '0563C1')
    u = etree.SubElement(rPr, qn('w:u'))
    u.set(qn('w:val'), 'single')
    if font_name:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    if font_size:
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), str(int(font_size * 2)))
        szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), str(int(font_size * 2)))
    t = etree.SubElement(run_elem, qn('w:t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = display_text or url


def _set_run_font(run, font_name=None, font_size=None):
    """Apply font name and size to a run. Always uses Symbol font for symbol characters."""
    if font_name:
        run.font.name = font_name
        # Ensure East Asian and complex script fonts also match
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    if font_size:
        run.font.size = Pt(font_size) if isinstance(font_size, (int, float)) else font_size


def _suppress_auto_numbering(p):
    """Add w:numPr with numId=0 to prevent Word from auto-converting
    paragraphs starting with '1. ' into a numbered list."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(p._element, qn('w:pPr'))
        p._element.insert(0, pPr)
    numPr = etree.SubElement(pPr, qn('w:numPr'))
    numId = etree.SubElement(numPr, qn('w:numId'))
    numId.set(qn('w:val'), '0')


def _add_formatted_runs(p, text, font_name=None, font_size=None,
                        superscript_citations=False):
    """Add runs to paragraph p, handling citations and markdown formatting."""
    last = 0
    for m in _FMT_RE.finditer(text):
        # Add any preceding plain text
        if m.start() > last:
            plain = text[last:m.start()]
            if plain:
                r = p.add_run(plain)
                _set_run_font(r, font_name, font_size)
        cite, bolditalic, bold, italic, subscr, url = m.group(1), m.group(2), m.group(3), m.group(4), m.group(5), m.group(6)
        if url:
            _add_hyperlink(p, url, font_name=font_name, font_size=font_size)
        elif cite and superscript_citations:
            r = p.add_run(cite[1:-1])  # strip [ ]
            r.font.superscript = True
            _set_run_font(r, font_name, font_size)
        elif bolditalic:
            inner = bolditalic[3:-3]
            r = p.add_run(inner)
            r.bold = True
            r.italic = True
            _set_run_font(r, font_name, font_size)
        elif bold:
            inner = bold[2:-2]
            r = p.add_run(inner)
            r.bold = True
            _set_run_font(r, font_name, font_size)
        elif italic:
            inner = italic[1:-1]
            r = p.add_run(inner)
            r.italic = True
            _set_run_font(r, font_name, font_size)
        elif subscr:
            base, sub = subscr.split('_', 1)
            r = p.add_run(base)
            _set_run_font(r, font_name, font_size)
            r = p.add_run(sub)
            r.font.subscript = True
            _set_run_font(r, font_name, font_size)
        else:
            r = p.add_run(m.group(0))
            _set_run_font(r, font_name, font_size)
        last = m.end()
    # Remaining text
    if last < len(text):
        r = p.add_run(text[last:])
        _set_run_font(r, font_name, font_size)


def add_body_paragraphs(doc, text, font_name=None, font_size=None,
                        is_references=False, superscript_citations=False,
                        suppress_first_indent=False):
    """Add body text, splitting on double newlines into separate paragraphs.

    If is_references=True, each paragraph is formatted with left alignment
    and a hanging indent (first line outdented relative to the rest).
    References are split on single newlines (each reference = one line).
    If suppress_first_indent=True, the first paragraph has no first-line indent
    (used after headings). All paragraphs use Normal style.
    """
    if is_references:
        # References are one-per-line; split on single newlines
        paragraphs = text.split('\n')
    else:
        paragraphs = text.split('\n\n')
    total_words = 0
    for i, para_text in enumerate(paragraphs):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(style='Normal')
        _add_formatted_runs(p, para_text, font_name, font_size,
                            superscript_citations=superscript_citations)
        if is_references:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            # Suppress Word auto-numbering so "1. " stays as literal text
            _suppress_auto_numbering(p)
        elif i == 0 and suppress_first_indent:
            p.paragraph_format.first_line_indent = Pt(0)
        total_words += len(para_text.split())
    return total_words


_INSERT_FIG_RE = re.compile(r'\[INSERT FIGURE (\d+) HERE\]')


def _add_table_caption(doc, caption_text, font_name=None, font_size=None,
                       superscript_citations=False):
    """Add a table caption in Subtitle style: 'Table N.' bold, rest normal weight.

    Formatting matches reference doc: Subtitle style (11pt, inherits from Caption),
    single-spaced, space_after=6pt, keep_with_next=True.
    Supports inline $..$ math in caption text.
    """
    try:
        p = doc.add_paragraph(style='Subtitle')
    except KeyError:
        p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    # Parse "**Table N.** rest of caption" or "Table N. rest"
    m = re.match(r'\*\*(.+?)\*\*\s*(.*)', caption_text)
    if m:
        r = p.add_run(m.group(1))
        r.bold = True
        _set_run_font(r, font_name, font_size)
        rest = ' ' + m.group(2) if m.group(2) else ''
        if rest.strip():
            # Handle inline math in caption
            if '$' in rest:
                parts = _INLINE_MATH_RE.split(rest)
                for idx, part in enumerate(parts):
                    if idx % 2 == 0:
                        if part:
                            r = p.add_run(part)
                            r.bold = False
                            _set_run_font(r, font_name, font_size)
                    else:
                        _add_omml_equation(p, part, font_name)
            else:
                r = p.add_run(rest)
                r.bold = False
                _set_run_font(r, font_name, font_size)
    else:
        _add_formatted_runs(p, caption_text, font_name, font_size,
                            superscript_citations=superscript_citations)
    return p


def _add_figure_caption(doc, caption_text, font_name=None, font_size=10,
                        superscript_citations=False):
    """Add a figure caption in Caption style.

    "Figure N." prefix is bold, rest is normal weight. Caption style has
    1.5 line spacing (set in template). Supports inline $..$ math.
    """
    try:
        p = doc.add_paragraph(style='Caption')
    except KeyError:
        p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Pt(0)
    # "Figure N." should be bold
    m = re.match(r'(Figure [A-Za-z]?\d+\.)\s*(.*)', caption_text, re.DOTALL)
    if m:
        r = p.add_run(m.group(1))
        r.bold = True
        _set_run_font(r, font_name, font_size)
        rest = ' ' + m.group(2) if m.group(2) else ''
        if rest.strip():
            # Handle inline math in caption
            if '$' in rest:
                parts = _INLINE_MATH_RE.split(rest)
                for idx, part in enumerate(parts):
                    if idx % 2 == 0:
                        if part:
                            _add_formatted_runs(p, part, font_name, font_size,
                                                superscript_citations=superscript_citations)
                    else:
                        _add_omml_equation(p, part, font_name)
            else:
                _add_formatted_runs(p, rest, font_name, font_size,
                                    superscript_citations=superscript_citations)
    else:
        _add_formatted_runs(p, caption_text, font_name, font_size,
                            superscript_citations=superscript_citations)
    return p


def _add_table_note(doc, note_text, font_name=None, font_size=10,
                    superscript_citations=False):
    """Add a table note in Caption style, prefixed with bold 'Note. '.

    Same style as figure captions (Caption, 1.5 line spacing from template).
    Supports inline $..$ math in note text.
    """
    try:
        p = doc.add_paragraph(style='Caption')
    except KeyError:
        p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Pt(0)
    # Strip any existing Note prefix (plain, italic, or bold markdown) and add bold one
    note_text = re.sub(r'^\*{0,2}Notes?[\.:]\*{0,2}\s*', '', note_text)
    r = p.add_run('Note. ')
    r.bold = True
    _set_run_font(r, font_name, font_size)
    # Handle inline math in note text
    if '$' in note_text:
        parts = _INLINE_MATH_RE.split(note_text)
        for idx, part in enumerate(parts):
            if idx % 2 == 0:
                if part:
                    _add_formatted_runs(p, part, font_name, font_size,
                                        superscript_citations=superscript_citations)
            else:
                _add_omml_equation(p, part, font_name)
    else:
        _add_formatted_runs(p, note_text, font_name, font_size,
                            superscript_citations=superscript_citations)
    return p


def _add_cell_content(p, text, font_name=None, font_size=9, is_header=False):
    """Add content to a table cell paragraph, handling inline math and markdown.

    Detects cell-level bold (entire cell wrapped in **...**) or bold+italic
    (***...***), then splits on $..$ for inline math (OMML), and uses
    _add_formatted_runs for remaining markdown formatting.
    """
    text = text.strip()
    if not text:
        return
    # Detect cell-level bold/italic wrapping (** or *** around entire cell)
    cell_bold = False
    cell_italic = False
    if text.startswith('***') and text.endswith('***'):
        text = text[3:-3]
        cell_bold = True
        cell_italic = True
    elif text.startswith('**') and text.endswith('**'):
        text = text[2:-2]
        cell_bold = True

    if '$' in text:
        parts = _INLINE_MATH_RE.split(text)
        for idx, part in enumerate(parts):
            if idx % 2 == 0:
                if part:
                    _add_formatted_runs(p, part, font_name, font_size)
            else:
                _add_omml_equation(p, part, font_name)
    else:
        _add_formatted_runs(p, text, font_name, font_size)

    # Apply cell-level formatting to all runs
    if cell_bold or cell_italic or is_header:
        for run in p.runs:
            if cell_bold or is_header:
                run.bold = True
            if cell_italic:
                run.italic = True


def _add_table_from_tsv(doc, tsv_text, font_name=None, cell_font_size=11):
    """Parse tab-separated text into a proper Word table.

    Plain Table 1 style, centered in document, autofit to contents,
    11pt font, centered cells, single-spaced, space_before/after=1pt,
    header row bold. Supports inline $..$ math and *italic*/**bold** markdown.
    """
    lines = [l for l in tsv_text.strip().split('\n') if l.strip()]
    if not lines:
        return
    rows_data = [line.split('\t') for line in lines]
    n_cols = max(len(r) for r in rows_data)
    # Pad short rows
    for r in rows_data:
        while len(r) < n_cols:
            r.append('')

    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    # Apply Plain Table 1 style (falls back to Table Grid if unavailable)
    for style_name in ('Plain Table 1', 'Table Grid'):
        try:
            table.style = doc.styles[style_name]
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Autofit to contents
    tbl_pr = table._tbl.tblPr
    # Set table width to auto
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = etree.SubElement(tbl_pr, qn('w:tblW'))
    tbl_w.set(qn('w:type'), 'auto')
    tbl_w.set(qn('w:w'), '0')
    # Remove fixed layout (enables autofit)
    tbl_layout = tbl_pr.find(qn('w:tblLayout'))
    if tbl_layout is not None:
        tbl_pr.remove(tbl_layout)
    # Set all cell widths to auto so Word autofits to content
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = etree.SubElement(tc_pr, qn('w:tcW'))
            tc_w.set(qn('w:type'), 'auto')
            tc_w.set(qn('w:w'), '0')

    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            # Clear default paragraph
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            # 1pt space before/after, single-spaced
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = Pt(0)
            # Add content with math + formatting support
            _add_cell_content(p, cell_text, font_name, cell_font_size,
                              is_header=(ri == 0))

    # Remove all color/shading from cells and set keep_with_next on all but last row
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn('w:shd'))
            if shading is not None:
                tc_pr.remove(shading)
            if ri < len(table.rows) - 1:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True

    return table


def _find_figure_file(fig_number, figures_dir):
    """Find a figure file by number, trying common extensions.

    fig_number can be numeric ('1', '2') or alphanumeric ('S1', 'S2', 'A1').
    Searches for patterns like figure1.png, figure_s1_*.png, Figure_S1.png, etc.
    """
    import glob
    # Normalize: try both original and lowercase
    num = str(fig_number)
    num_lower = num.lower()
    # Try exact patterns first
    for pattern_template in [
        'Figure{num}.{ext}', 'figure{num}.{ext}',
        'Fig{num}.{ext}', 'fig{num}.{ext}',
        'Figure_{num}.{ext}', 'figure_{num}.{ext}',
        'Figure-{num}.{ext}', 'figure-{num}.{ext}',
    ]:
        for ext in ['png', 'tif', 'tiff', 'jpg', 'jpeg']:
            pattern = os.path.join(figures_dir,
                                   pattern_template.format(num=num, ext=ext))
            matches = glob.glob(pattern)
            if matches:
                return matches[0]
            # Try lowercase variant
            if num != num_lower:
                pattern = os.path.join(figures_dir,
                                       pattern_template.format(num=num_lower, ext=ext))
                matches = glob.glob(pattern)
                if matches:
                    return matches[0]
    # Try wildcard patterns (e.g., figure_s1_endorsement_merged.png)
    for prefix in ['figure', 'Figure', 'fig', 'Fig']:
        for sep in ['_', '-', '']:
            for ext in ['png', 'tif', 'tiff', 'jpg', 'jpeg']:
                pattern = os.path.join(figures_dir,
                                       f'{prefix}{sep}{num_lower}*{ext}')
                matches = glob.glob(pattern)
                if matches:
                    return matches[0]
                # Also try with underscore after number
                pattern = os.path.join(figures_dir,
                                       f'{prefix}{sep}{num_lower}_*.{ext}')
                matches = glob.glob(pattern)
                if matches:
                    return matches[0]
    return None


def _insert_figure(doc, fig_number, figures_dir, font_name=None, font_size=None,
                   fig_path=None):
    """Insert a figure image inline in a centered paragraph.

    Image is inserted at page width (6.5in for 1in margins) in a centered
    Normal paragraph with keep_with_next=True so it stays with its caption.

    If ``fig_path`` is provided (explicit path from MD ``![](path)`` or caller),
    it is used as-is; otherwise the file is located by scanning ``figures_dir``
    for ``Figure{N}.png`` and similar filename patterns.
    """
    if not fig_path or not os.path.isfile(fig_path):
        fig_path = _find_figure_file(fig_number, figures_dir)
    if not fig_path:
        return False

    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True  # Keep with caption below
    p.paragraph_format.keep_together = True   # Never split image across pages
    r = p.add_run()
    r.add_picture(fig_path, width=Inches(6.5))
    return True


def _insert_grouped_figure(doc, fig_number, caption_text, figures_dir,
                           font_name=None, font_size=None,
                           superscript_citations=False,
                           target_width_inches=6.5,
                           page_width_inches=6.5,
                           fig_path=None):
    """Insert a figure with caption as a grouped inline object.

    Creates an inline DrawingML group (wpg:wgp) containing:
    1. pic:pic — the image (centered within the group)
    2. wps:wsp — a text box with the caption in Caption style

    The group is always full page width so the caption spans the whole line,
    even when the image is narrower. The group uses inline (In Line with Text)
    wrapping, matching Word's behavior when you group objects.

    If ``fig_path`` is provided (explicit path from MD ``![](path)``), it is
    used as-is; otherwise the file is located by scanning ``figures_dir``.
    """
    from PIL import Image as PILImage

    if not fig_path or not os.path.isfile(fig_path):
        fig_path = _find_figure_file(fig_number, figures_dir)
    if not fig_path:
        return False

    # Get image dimensions
    img = PILImage.open(fig_path)
    img_w_px, img_h_px = img.size
    img.close()

    # Image scales to target_width_inches (may be narrower than page)
    target_w_emu = Inches(target_width_inches)
    scale = target_w_emu / Emu(int(img_w_px * 914400 / 96))  # 96 dpi assumed
    img_w_emu = target_w_emu
    img_h_emu = Emu(int(img_h_px * 914400 / 96 * scale))

    # Group is always full page width so caption spans the whole line
    group_w_emu = Inches(page_width_inches)

    # Caption text box height estimate (generous)
    caption_h_emu = Emu(int(Inches(1.5)))
    group_h_emu = img_h_emu + caption_h_emu

    # Add the image to the document's media relationships
    rId, _image = doc.part.get_or_add_image(fig_path)

    # Get unique IDs for the drawing objects
    import random
    base_id = random.randint(100000, 999999)

    # Namespace declarations
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    WPG_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup'
    WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    MC_NS = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    WP14_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing'
    W_NS = _W_NS

    nsmap = {
        'wp': WP_NS, 'a': A_NS, 'pic': PIC_NS, 'r': R_NS,
        'wpg': WPG_NS, 'wps': WPS_NS, 'mc': MC_NS, 'wp14': WP14_NS,
        'w': W_NS,
    }

    # Build the DrawingML XML
    # Create paragraph (no extra spacing around figure)
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    # w:r > mc:AlternateContent > mc:Choice Requires="wpg" > w:drawing > wp:inline
    # This matches the exact XML structure Word produces when you group
    # objects and set them to "In Line with Text" wrapping.
    run_elem = etree.SubElement(p._element, f'{{{W_NS}}}r')
    # noProof on the run (Word adds this for drawing runs)
    run_rPr = etree.SubElement(run_elem, f'{{{W_NS}}}rPr')
    etree.SubElement(run_rPr, f'{{{W_NS}}}noProof')
    mc_alt = etree.SubElement(run_elem, f'{{{MC_NS}}}AlternateContent')
    mc_choice = etree.SubElement(mc_alt, f'{{{MC_NS}}}Choice')
    mc_choice.set('Requires', 'wpg')
    drawing = etree.SubElement(mc_choice, f'{{{W_NS}}}drawing')
    inline = etree.SubElement(drawing, f'{{{WP_NS}}}inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')

    # Extent
    extent = etree.SubElement(inline, f'{{{WP_NS}}}extent')
    extent.set('cx', str(int(group_w_emu)))
    extent.set('cy', str(int(group_h_emu)))

    # effectExtent
    effectExtent = etree.SubElement(inline, f'{{{WP_NS}}}effectExtent')
    effectExtent.set('l', '0')
    effectExtent.set('t', '0')
    effectExtent.set('r', '0')
    effectExtent.set('b', '0')

    # docPr
    docPr = etree.SubElement(inline, f'{{{WP_NS}}}docPr')
    docPr.set('id', str(base_id))
    docPr.set('name', f'Group {fig_number}')

    # cNvGraphicFramePr
    etree.SubElement(inline, f'{{{WP_NS}}}cNvGraphicFramePr')

    # a:graphic > a:graphicData > wpg:wgp
    graphic = etree.SubElement(inline, f'{{{A_NS}}}graphic')
    graphicData = etree.SubElement(graphic, f'{{{A_NS}}}graphicData')
    graphicData.set('uri', 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup')

    wgp = etree.SubElement(graphicData, f'{{{WPG_NS}}}wgp')

    # wpg:cNvGrpSpPr
    etree.SubElement(wgp, f'{{{WPG_NS}}}cNvGrpSpPr')

    # wpg:grpSpPr — group transform
    grpSpPr = etree.SubElement(wgp, f'{{{WPG_NS}}}grpSpPr')
    xfrm = etree.SubElement(grpSpPr, f'{{{A_NS}}}xfrm')
    off = etree.SubElement(xfrm, f'{{{A_NS}}}off')
    off.set('x', '0')
    off.set('y', '0')
    ext = etree.SubElement(xfrm, f'{{{A_NS}}}ext')
    ext.set('cx', str(int(group_w_emu)))
    ext.set('cy', str(int(group_h_emu)))
    chOff = etree.SubElement(xfrm, f'{{{A_NS}}}chOff')
    chOff.set('x', '0')
    chOff.set('y', '0')
    chExt = etree.SubElement(xfrm, f'{{{A_NS}}}chExt')
    chExt.set('cx', str(int(group_w_emu)))
    chExt.set('cy', str(int(group_h_emu)))

    # --- Child 1: pic:pic (the image) ---
    pic = etree.SubElement(wgp, f'{{{PIC_NS}}}pic')
    nvPicPr = etree.SubElement(pic, f'{{{PIC_NS}}}nvPicPr')
    cNvPr = etree.SubElement(nvPicPr, f'{{{PIC_NS}}}cNvPr')
    cNvPr.set('id', str(base_id + 1))
    cNvPr.set('name', f'Picture {fig_number}')
    cNvPicPr = etree.SubElement(nvPicPr, f'{{{PIC_NS}}}cNvPicPr')

    blipFill = etree.SubElement(pic, f'{{{PIC_NS}}}blipFill')
    blip = etree.SubElement(blipFill, f'{{{A_NS}}}blip')
    blip.set(f'{{{R_NS}}}embed', rId)
    stretch = etree.SubElement(blipFill, f'{{{A_NS}}}stretch')
    etree.SubElement(stretch, f'{{{A_NS}}}fillRect')

    spPr_pic = etree.SubElement(pic, f'{{{PIC_NS}}}spPr')
    xfrm_pic = etree.SubElement(spPr_pic, f'{{{A_NS}}}xfrm')
    off_pic = etree.SubElement(xfrm_pic, f'{{{A_NS}}}off')
    # Center the image horizontally within the group
    img_offset_x = max(0, (int(group_w_emu) - int(img_w_emu)) // 2)
    off_pic.set('x', str(img_offset_x))
    off_pic.set('y', '0')
    ext_pic = etree.SubElement(xfrm_pic, f'{{{A_NS}}}ext')
    ext_pic.set('cx', str(int(img_w_emu)))
    ext_pic.set('cy', str(int(img_h_emu)))
    prstGeom_pic = etree.SubElement(spPr_pic, f'{{{A_NS}}}prstGeom')
    prstGeom_pic.set('prst', 'rect')
    etree.SubElement(prstGeom_pic, f'{{{A_NS}}}avLst')

    # --- Child 2: wps:wsp (text box with caption) ---
    wsp = etree.SubElement(wgp, f'{{{WPS_NS}}}wsp')
    cNvPr2 = etree.SubElement(wsp, f'{{{WPS_NS}}}cNvPr')
    cNvPr2.set('id', str(base_id + 2))
    cNvPr2.set('name', f'Text Box {fig_number}')
    cNvSpPr = etree.SubElement(wsp, f'{{{WPS_NS}}}cNvSpPr')
    cNvSpPr.set('txBox', '1')

    spPr_wsp = etree.SubElement(wsp, f'{{{WPS_NS}}}spPr')
    xfrm_wsp = etree.SubElement(spPr_wsp, f'{{{A_NS}}}xfrm')
    off_wsp = etree.SubElement(xfrm_wsp, f'{{{A_NS}}}off')
    off_wsp.set('x', '0')
    off_wsp.set('y', str(int(img_h_emu)))
    ext_wsp = etree.SubElement(xfrm_wsp, f'{{{A_NS}}}ext')
    ext_wsp.set('cx', str(int(group_w_emu)))
    ext_wsp.set('cy', str(int(caption_h_emu)))
    prstGeom_wsp = etree.SubElement(spPr_wsp, f'{{{A_NS}}}prstGeom')
    prstGeom_wsp.set('prst', 'rect')
    etree.SubElement(prstGeom_wsp, f'{{{A_NS}}}avLst')
    # White fill
    solidFill = etree.SubElement(spPr_wsp, f'{{{A_NS}}}solidFill')
    srgbClr = etree.SubElement(solidFill, f'{{{A_NS}}}srgbClr')
    srgbClr.set('val', 'FFFFFF')
    # No border
    ln = etree.SubElement(spPr_wsp, f'{{{A_NS}}}ln')
    etree.SubElement(ln, f'{{{A_NS}}}noFill')

    # Text box content
    txbx = etree.SubElement(wsp, f'{{{WPS_NS}}}txbx')
    txbxContent = etree.SubElement(txbx, f'{{{W_NS}}}txbxContent')

    # Caption paragraph inside text box
    cap_p = etree.SubElement(txbxContent, f'{{{W_NS}}}p')
    cap_pPr = etree.SubElement(cap_p, f'{{{W_NS}}}pPr')
    pStyle = etree.SubElement(cap_pPr, f'{{{W_NS}}}pStyle')
    pStyle.set(f'{{{W_NS}}}val', 'Caption')

    # Parse caption: "Figure N." bold, rest normal
    cap_match = re.match(r'(Figure [A-Za-z]?\d+\.)\s*(.*)', caption_text, re.DOTALL)
    if cap_match:
        # Bold prefix
        r1 = etree.SubElement(cap_p, f'{{{W_NS}}}r')
        r1Pr = etree.SubElement(r1, f'{{{W_NS}}}rPr')
        etree.SubElement(r1Pr, f'{{{W_NS}}}b')
        if font_name:
            rFonts = etree.SubElement(r1Pr, f'{{{W_NS}}}rFonts')
            rFonts.set(f'{{{W_NS}}}ascii', font_name)
            rFonts.set(f'{{{W_NS}}}hAnsi', font_name)
        r1_t = etree.SubElement(r1, f'{{{W_NS}}}t')
        r1_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r1_t.text = cap_match.group(1) + ' '

        # Rest of caption — handle inline math ($...$) and **bold**/*italic*
        # markdown. Helper builds a w:r with optional bold/italic.
        def _emit_xml_run(parent, txt, bold_flag=False, italic_flag=False):
            if not txt:
                return
            r = etree.SubElement(parent, f'{{{W_NS}}}r')
            rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
            if bold_flag:
                etree.SubElement(rPr, f'{{{W_NS}}}b')
            if italic_flag:
                etree.SubElement(rPr, f'{{{W_NS}}}i')
            if font_name:
                rFonts = etree.SubElement(rPr, f'{{{W_NS}}}rFonts')
                rFonts.set(f'{{{W_NS}}}ascii', font_name)
                rFonts.set(f'{{{W_NS}}}hAnsi', font_name)
            t = etree.SubElement(r, f'{{{W_NS}}}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = txt

        def _emit_with_markdown(parent, txt):
            """Walk txt segmenting on ***bold-italic***, **bold**, *italic*."""
            md_re = re.compile(r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)')
            cursor = 0
            for m in md_re.finditer(txt):
                if m.start() > cursor:
                    _emit_xml_run(parent, txt[cursor:m.start()])
                token = m.group(0)
                if token.startswith('***'):
                    _emit_xml_run(parent, token[3:-3], bold_flag=True, italic_flag=True)
                elif token.startswith('**'):
                    _emit_xml_run(parent, token[2:-2], bold_flag=True)
                else:
                    _emit_xml_run(parent, token[1:-1], italic_flag=True)
                cursor = m.end()
            if cursor < len(txt):
                _emit_xml_run(parent, txt[cursor:])

        rest = cap_match.group(2)
        if rest:
            if '$' in rest:
                parts = _INLINE_MATH_RE.split(rest)
                for idx, part in enumerate(parts):
                    if idx % 2 == 0:
                        if part:
                            _emit_with_markdown(cap_p, part)
                    else:
                        # Build OMML equation and append to paragraph
                        tokens = _parse_latex_to_tokens(part)
                        oMath = etree.Element(f'{{{_OMML_NS}}}oMath')
                        for tok in tokens:
                            oMath.append(_build_element_from_token(tok, font_name))
                        cap_p.append(oMath)
            else:
                _emit_with_markdown(cap_p, rest)
    else:
        r1 = etree.SubElement(cap_p, f'{{{W_NS}}}r')
        r1_t = etree.SubElement(r1, f'{{{W_NS}}}t')
        r1_t.text = caption_text

    # Text body properties (zero insets, auto-resize shape to fit text)
    bodyPr = etree.SubElement(wsp, f'{{{WPS_NS}}}bodyPr')
    bodyPr.set('lIns', '0')
    bodyPr.set('tIns', '0')
    bodyPr.set('rIns', '0')
    bodyPr.set('bIns', '0')
    etree.SubElement(bodyPr, f'{{{A_NS}}}spAutoFit')

    # mc:Fallback — inline image for processors that don't support wpg groups
    mc_fallback = etree.SubElement(mc_alt, f'{{{MC_NS}}}Fallback')
    fb_drawing = etree.SubElement(mc_fallback, f'{{{W_NS}}}drawing')
    fb_inline = etree.SubElement(fb_drawing, f'{{{WP_NS}}}inline')
    fb_inline.set('distT', '0')
    fb_inline.set('distB', '0')
    fb_inline.set('distL', '0')
    fb_inline.set('distR', '0')
    fb_extent = etree.SubElement(fb_inline, f'{{{WP_NS}}}extent')
    fb_extent.set('cx', str(int(img_w_emu)))
    fb_extent.set('cy', str(int(img_h_emu)))
    fb_effectExtent = etree.SubElement(fb_inline, f'{{{WP_NS}}}effectExtent')
    fb_effectExtent.set('l', '0')
    fb_effectExtent.set('t', '0')
    fb_effectExtent.set('r', '0')
    fb_effectExtent.set('b', '0')
    fb_docPr = etree.SubElement(fb_inline, f'{{{WP_NS}}}docPr')
    fb_docPr.set('id', str(base_id + 10))
    fb_docPr.set('name', f'Picture {fig_number} Fallback')
    etree.SubElement(fb_inline, f'{{{WP_NS}}}cNvGraphicFramePr')
    fb_graphic = etree.SubElement(fb_inline, f'{{{A_NS}}}graphic')
    fb_graphicData = etree.SubElement(fb_graphic, f'{{{A_NS}}}graphicData')
    fb_graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    fb_pic = etree.SubElement(fb_graphicData, f'{{{PIC_NS}}}pic')
    fb_nvPicPr = etree.SubElement(fb_pic, f'{{{PIC_NS}}}nvPicPr')
    fb_cNvPr = etree.SubElement(fb_nvPicPr, f'{{{PIC_NS}}}cNvPr')
    fb_cNvPr.set('id', str(base_id + 11))
    fb_cNvPr.set('name', f'Picture {fig_number}')
    etree.SubElement(fb_nvPicPr, f'{{{PIC_NS}}}cNvPicPr')
    fb_blipFill = etree.SubElement(fb_pic, f'{{{PIC_NS}}}blipFill')
    fb_blip = etree.SubElement(fb_blipFill, f'{{{A_NS}}}blip')
    fb_blip.set(f'{{{R_NS}}}embed', rId)
    fb_stretch = etree.SubElement(fb_blipFill, f'{{{A_NS}}}stretch')
    etree.SubElement(fb_stretch, f'{{{A_NS}}}fillRect')
    fb_spPr = etree.SubElement(fb_pic, f'{{{PIC_NS}}}spPr')
    fb_xfrm = etree.SubElement(fb_spPr, f'{{{A_NS}}}xfrm')
    fb_off = etree.SubElement(fb_xfrm, f'{{{A_NS}}}off')
    fb_off.set('x', '0')
    fb_off.set('y', '0')
    fb_ext = etree.SubElement(fb_xfrm, f'{{{A_NS}}}ext')
    fb_ext.set('cx', str(int(img_w_emu)))
    fb_ext.set('cy', str(int(img_h_emu)))
    fb_prstGeom = etree.SubElement(fb_spPr, f'{{{A_NS}}}prstGeom')
    fb_prstGeom.set('prst', 'rect')
    etree.SubElement(fb_prstGeom, f'{{{A_NS}}}avLst')

    return True


# ── OMML equation support ────────────────────────────────────────────

_OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _omml_run(text, font_name=None, italic=True):
    """Create an OMML run (<m:r>) element with the given text."""
    r = etree.SubElement(etree.Element('tmp'), f'{{{_OMML_NS}}}r')
    # Run properties
    rPr = etree.SubElement(r, f'{{{_OMML_NS}}}rPr')
    sty = etree.SubElement(rPr, f'{{{_OMML_NS}}}sty')
    sty.set(f'{{{_OMML_NS}}}val', 'i' if italic else 'p')
    # Word run properties (font)
    if font_name:
        w_rPr = etree.SubElement(rPr, f'{{{_W_NS}}}rPr')
        rFonts = etree.SubElement(w_rPr, f'{{{_W_NS}}}rFonts')
        rFonts.set(f'{{{_W_NS}}}ascii', font_name)
        rFonts.set(f'{{{_W_NS}}}hAnsi', font_name)
    # Text
    t = etree.SubElement(r, f'{{{_OMML_NS}}}t')
    t.text = text
    return r


def _omml_sub_elem(base_elem, sub_elem):
    """Create an OMML subscript from pre-built elements."""
    sSub = etree.Element(f'{{{_OMML_NS}}}sSub')
    e_base = etree.SubElement(sSub, f'{{{_OMML_NS}}}e')
    e_base.append(base_elem)
    e_sub = etree.SubElement(sSub, f'{{{_OMML_NS}}}sub')
    e_sub.append(sub_elem)
    return sSub


def _omml_sub(base_text, sub_text, font_name=None):
    """Create an OMML subscript element (<m:sSub>)."""
    return _omml_sub_elem(
        _omml_run(base_text, font_name),
        _omml_run(sub_text, font_name))


def _omml_sup_elem(base_elem, sup_elem):
    """Create an OMML superscript from pre-built elements."""
    sSup = etree.Element(f'{{{_OMML_NS}}}sSup')
    e_base = etree.SubElement(sSup, f'{{{_OMML_NS}}}e')
    e_base.append(base_elem)
    e_sup = etree.SubElement(sSup, f'{{{_OMML_NS}}}sup')
    e_sup.append(sup_elem)
    return sSup


def _omml_sup(base_text, sup_text, font_name=None):
    """Create an OMML superscript element (<m:sSup>)."""
    return _omml_sup_elem(
        _omml_run(base_text, font_name),
        _omml_run(sup_text, font_name))


def _omml_subsup_elem(base_elem, sub_elem, sup_elem):
    """Create an OMML simultaneous sub+superscript element (<m:sSubSup>)."""
    sSubSup = etree.Element(f'{{{_OMML_NS}}}sSubSup')
    e_base = etree.SubElement(sSubSup, f'{{{_OMML_NS}}}e')
    e_base.append(base_elem)
    e_sub = etree.SubElement(sSubSup, f'{{{_OMML_NS}}}sub')
    e_sub.append(sub_elem)
    e_sup = etree.SubElement(sSubSup, f'{{{_OMML_NS}}}sup')
    e_sup.append(sup_elem)
    return sSubSup


def _omml_accent(base_elem, accent_char):
    r"""Create an OMML accent element (<m:acc>) for \bar, \hat, \tilde, etc."""
    acc = etree.Element(f'{{{_OMML_NS}}}acc')
    accPr = etree.SubElement(acc, f'{{{_OMML_NS}}}accPr')
    chrElem = etree.SubElement(accPr, f'{{{_OMML_NS}}}chr')
    chrElem.set(f'{{{_OMML_NS}}}val', accent_char)
    e = etree.SubElement(acc, f'{{{_OMML_NS}}}e')
    e.append(base_elem)
    return acc


def _token_to_element(tok, font_name=None):
    """Convert a single parsed token to an OMML element."""
    if tok['type'] == 'text':
        return _omml_run(tok['text'], font_name, italic=True)
    elif tok['type'] == 'plain':
        return _omml_run(tok['text'], font_name, italic=False)
    elif tok['type'] == 'greek':
        char = _GREEK_MAP.get(tok['text'], tok['text'])
        return _omml_run(char, font_name, italic=True)
    elif tok['type'] == 'symbol':
        char = _MATH_SYMBOLS.get(tok['text'], tok['text'])
        return _omml_run(char, font_name, italic=False)
    elif tok['type'] == 'sub':
        base_elem = tok.get('base_elem')
        sub_elem = tok.get('sub_elem')
        if base_elem is not None and sub_elem is not None:
            return _omml_sub_elem(base_elem, sub_elem)
        base = tok.get('base', '?')
        sub = tok.get('sub', '')
        return _omml_sub(base, sub, font_name)
    elif tok['type'] == 'sup':
        base_elem = tok.get('base_elem')
        sup_elem = tok.get('sup_elem')
        if base_elem is not None and sup_elem is not None:
            return _omml_sup_elem(base_elem, sup_elem)
        base = tok.get('base', '?')
        sup = tok.get('sup', '')
        return _omml_sup(base, sup, font_name)
    elif tok['type'] == 'accent':
        base_elem = tok.get('base_elem')
        if base_elem is None:
            base_elem = _omml_run(tok.get('base', '?'), font_name)
        return _omml_accent(base_elem, tok['char'])
    else:
        return _omml_run(tok.get('text', '?'), font_name, italic=True)


def _build_omml_from_tokens(tokens, font_name=None):
    """Build OMML math zone (<m:oMath>) from parsed tokens."""
    oMath = etree.Element(f'{{{_OMML_NS}}}oMath')
    for tok in tokens:
        oMath.append(_token_to_element(tok, font_name))
    return oMath


def _read_brace_arg(s, i):
    """Read a brace-delimited argument {content} starting at position i.
    Returns (content, new_position). If no brace, reads single char."""
    if i >= len(s):
        return '', i
    if s[i] == '{':
        depth = 1
        j = i + 1
        while j < len(s) and depth > 0:
            if s[j] == '{':
                depth += 1
            elif s[j] == '}':
                depth -= 1
            j += 1
        return s[i+1:j-1], j
    else:
        return s[i], i + 1


def _parse_latex_to_tokens(latex_str):
    """Parse a LaTeX math string into tokens for OMML conversion.

    Handles: \\greek, \\symbol, \\text{}, \\bar{}, \\hat{},
    \\operatorname{}, x_{sub}, x^{sup}, \\,  spacing, and nesting.
    """
    tokens = []
    i = 0
    s = latex_str.strip()
    while i < len(s):
        # Skip whitespace — equation spacing is handled by the renderer
        if s[i] == ' ':
            i += 1
            continue
        # LaTeX command: \commandname or \, \; \! etc.
        if s[i] == '\\':
            if i + 1 >= len(s):
                i += 1
                continue
            # Spacing commands: \, \; \! \: — all skipped (no extra spacing in equations)
            if s[i+1] in ',;!: ':
                i += 2
                continue
            # Named command
            if s[i+1].isalpha():
                j = i + 1
                while j < len(s) and s[j].isalpha():
                    j += 1
                cmd = s[i+1:j]
                i = j
                # \text{...} — upright text
                if cmd == 'text':
                    arg, i = _read_brace_arg(s, i)
                    tokens.append({'type': 'plain', 'text': arg})
                    continue
                # \mathrm{...} — same as \text
                if cmd == 'mathrm':
                    arg, i = _read_brace_arg(s, i)
                    tokens.append({'type': 'plain', 'text': arg})
                    continue
                # \mathbf{...} — bold math (render as italic for now)
                if cmd == 'mathbf':
                    arg, i = _read_brace_arg(s, i)
                    tokens.append({'type': 'text', 'text': arg})
                    continue
                # Accent commands: \bar{}, \hat{}, \tilde{}, etc.
                if cmd in _ACCENT_MAP:
                    arg, i = _read_brace_arg(s, i)
                    # Parse the argument recursively for Greek etc.
                    inner_tokens = _parse_latex_to_tokens(arg)
                    if len(inner_tokens) == 1:
                        base_elem_fn = inner_tokens[0]
                    else:
                        # Multiple tokens inside accent — just use text
                        base_elem_fn = {'type': 'text', 'text': arg}
                    tokens.append({
                        'type': 'accent',
                        'char': _ACCENT_MAP[cmd],
                        'base_token': base_elem_fn,
                    })
                    continue
                # Greek letters
                if cmd in _GREEK_MAP:
                    tokens.append({'type': 'greek', 'text': cmd})
                    continue
                # Math symbols
                if cmd in _MATH_SYMBOLS:
                    tokens.append({'type': 'symbol', 'text': cmd})
                    continue
                # Math operators (min, max, log, etc.) — upright text
                if cmd in _MATH_OPERATORS:
                    tokens.append({'type': 'plain', 'text': cmd})
                    continue
                # Unknown command — render as upright text
                tokens.append({'type': 'plain', 'text': cmd})
                continue
            # \\ (line break in display math) — skip
            if s[i+1] == '\\':
                i += 2
                continue
            # Other escaped chars
            i += 2
            continue
        # Subscript: _{...} or _x
        if s[i] == '_' and tokens:
            last = tokens.pop()
            i += 1
            sub_content, i = _read_brace_arg(s, i)
            sub_tokens = _parse_latex_to_tokens(sub_content)
            # Check if followed immediately by ^{...} → subsup
            if i < len(s) and s[i] == '^':
                i += 1
                sup_content, i = _read_brace_arg(s, i)
                sup_tokens = _parse_latex_to_tokens(sup_content)
                tokens.append({
                    'type': 'subsup',
                    'base_token': last,
                    'sub_tokens': sub_tokens,
                    'sup_tokens': sup_tokens,
                })
            # Check if last was already a sup → convert to subsup
            elif last.get('type') == 'sup':
                tokens.append({
                    'type': 'subsup',
                    'base_token': last['base_token'],
                    'sub_tokens': sub_tokens,
                    'sup_tokens': last['sup_tokens'],
                })
            else:
                tokens.append({
                    'type': 'sub',
                    'base_token': last,
                    'sub_tokens': sub_tokens,
                })
            continue
        # Superscript: ^{...} or ^x
        if s[i] == '^' and tokens:
            last = tokens.pop()
            i += 1
            sup_content, i = _read_brace_arg(s, i)
            sup_tokens = _parse_latex_to_tokens(sup_content)
            # Check if followed immediately by _{...} → subsup
            if i < len(s) and s[i] == '_':
                i += 1
                sub_content, i = _read_brace_arg(s, i)
                sub_tokens = _parse_latex_to_tokens(sub_content)
                tokens.append({
                    'type': 'subsup',
                    'base_token': last,
                    'sub_tokens': sub_tokens,
                    'sup_tokens': sup_tokens,
                })
            # Check if last was already a sub → convert to subsup
            elif last.get('type') == 'sub':
                tokens.append({
                    'type': 'subsup',
                    'base_token': last['base_token'],
                    'sub_tokens': last['sub_tokens'],
                    'sup_tokens': sup_tokens,
                })
            else:
                tokens.append({
                    'type': 'sup',
                    'base_token': last,
                    'sup_tokens': sup_tokens,
                })
            continue
        # Braces (skip)
        if s[i] in '{}':
            i += 1
            continue
        # Plain character
        tokens.append({'type': 'text', 'text': s[i]})
        i += 1
    return tokens


def _build_element_from_token(tok, font_name=None):
    """Recursively build an OMML element from a parsed token (new format)."""
    if tok['type'] in ('text', 'plain', 'greek', 'symbol'):
        return _token_to_element(tok, font_name)
    elif tok['type'] == 'sub':
        base_tok = tok.get('base_token')
        sub_toks = tok.get('sub_tokens', [])
        if base_tok is not None:
            base_elem = _build_element_from_token(base_tok, font_name)
        else:
            base_elem = _omml_run(tok.get('base', '?'), font_name)
        # Build sub element(s)
        if len(sub_toks) == 1:
            sub_elem = _build_element_from_token(sub_toks[0], font_name)
        elif sub_toks:
            # Multiple tokens in subscript — concatenate as text
            sub_text = ''.join(_token_display_text(t) for t in sub_toks)
            sub_elem = _omml_run(sub_text, font_name)
        else:
            sub_elem = _omml_run(tok.get('sub', ''), font_name)
        return _omml_sub_elem(base_elem, sub_elem)
    elif tok['type'] == 'sup':
        base_tok = tok.get('base_token')
        sup_toks = tok.get('sup_tokens', [])
        if base_tok is not None:
            base_elem = _build_element_from_token(base_tok, font_name)
        else:
            base_elem = _omml_run(tok.get('base', '?'), font_name)
        if len(sup_toks) == 1:
            sup_elem = _build_element_from_token(sup_toks[0], font_name)
        elif sup_toks:
            sup_text = ''.join(_token_display_text(t) for t in sup_toks)
            sup_elem = _omml_run(sup_text, font_name)
        else:
            sup_elem = _omml_run(tok.get('sup', ''), font_name)
        return _omml_sup_elem(base_elem, sup_elem)
    elif tok['type'] == 'subsup':
        base_tok = tok.get('base_token')
        sub_toks = tok.get('sub_tokens', [])
        sup_toks = tok.get('sup_tokens', [])
        base_elem = _build_element_from_token(base_tok, font_name) if base_tok else _omml_run('?', font_name)
        if len(sub_toks) == 1:
            sub_elem = _build_element_from_token(sub_toks[0], font_name)
        elif sub_toks:
            sub_text = ''.join(_token_display_text(t) for t in sub_toks)
            sub_elem = _omml_run(sub_text, font_name)
        else:
            sub_elem = _omml_run('', font_name)
        if len(sup_toks) == 1:
            sup_elem = _build_element_from_token(sup_toks[0], font_name)
        elif sup_toks:
            sup_text = ''.join(_token_display_text(t) for t in sup_toks)
            sup_elem = _omml_run(sup_text, font_name)
        else:
            sup_elem = _omml_run('', font_name)
        return _omml_subsup_elem(base_elem, sub_elem, sup_elem)
    elif tok['type'] == 'accent':
        base_tok = tok.get('base_token')
        if base_tok is not None:
            base_elem = _build_element_from_token(base_tok, font_name)
        else:
            base_elem = _omml_run(tok.get('base', '?'), font_name)
        return _omml_accent(base_elem, tok['char'])
    else:
        return _token_to_element(tok, font_name)


def _token_display_text(tok):
    """Get display text from a token (for concatenating in sub/superscripts)."""
    if tok['type'] == 'greek':
        return _GREEK_MAP.get(tok['text'], tok['text'])
    elif tok['type'] == 'symbol':
        return _MATH_SYMBOLS.get(tok['text'], tok['text'])
    elif tok['type'] in ('text', 'plain'):
        return tok['text']
    return tok.get('text', '?')


_COMPLEX_LATEX_RE = re.compile(
    r'\\(?:frac|dfrac|sqrt|mathbf|begin|left|right|pmatrix|bmatrix|text\{|mathrm\{|tag\{|quad|qquad|operatorname)')


def _add_omml_equation(p, latex_str, font_name=None):
    """Parse a LaTeX string and insert it as an OMML inline equation in paragraph p.

    Uses latex2mathml + mml2omml for complex LaTeX (fractions, matrices, roots, etc.)
    and the simple tokenizer for basic inline math (Greek letters, sub/superscripts).
    """
    if _COMPLEX_LATEX_RE.search(latex_str):
        try:
            import sys, os
            skill_dir = os.path.dirname(os.path.abspath(__file__))
            if skill_dir not in sys.path:
                sys.path.insert(0, skill_dir)
            from mml2omml import latex_to_omml
            oMath = latex_to_omml(latex_str)
            if oMath is not None:
                p._element.append(oMath)
                return
        except Exception as e:
            print(f"WARNING: mml2omml failed for '{latex_str[:60]}...': {e}, "
                  f"falling back to simple tokenizer")
    # Fallback: simple tokenizer for basic math
    tokens = _parse_latex_to_tokens(latex_str)
    oMath = etree.Element(f'{{{_OMML_NS}}}oMath')
    for tok in tokens:
        oMath.append(_build_element_from_token(tok, font_name))
    p._element.append(oMath)


# ── Inline math detection pattern ────────────────────────────────────
# Matches $..$ (inline math) in text
_INLINE_MATH_RE = re.compile(r'\$([^$]+?)\$')
# Matches $$...$$ (display math) in text
_DISPLAY_MATH_RE = re.compile(r'\$\$(.+?)\$\$', re.DOTALL)


def _add_paragraph_with_math(doc, para_text, font_name=None, font_size=None,
                             superscript_citations=False, style='Normal'):
    """Add a paragraph that may contain inline $..$ math expressions.

    Text outside $..$ is rendered as normal runs (with formatting).
    Text inside $..$ is rendered as OMML inline equations using Word's
    native equation editor format.
    """
    p = doc.add_paragraph(style=style)

    # Split on inline math delimiters
    parts = _INLINE_MATH_RE.split(para_text)
    # parts alternates: [text, math, text, math, text, ...]
    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            # Normal text
            if part:
                _add_formatted_runs(p, part, font_name, font_size,
                                    superscript_citations=superscript_citations)
        else:
            # Math content
            _add_omml_equation(p, part, font_name)

    return p


def add_body_paragraphs_with_math(doc, text, font_name=None, font_size=None,
                                  is_references=False, superscript_citations=False,
                                  suppress_first_indent=False):
    """Add body text with inline math support.

    Display math ($$...$$) gets its own centered paragraph.
    Inline math ($...$) is rendered as OMML within the paragraph.
    """
    # First, handle display math by splitting on $$..$$
    display_parts = _DISPLAY_MATH_RE.split(text)
    total_words = 0

    para_counter = 0
    for idx, part in enumerate(display_parts):
        if idx % 2 == 1:
            # Display math — its own centered paragraph
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            _add_omml_equation(p, part.strip(), font_name)
            total_words += len(part.split())
            continue

        # Normal text (may contain inline math and newline-separated paragraphs)
        if is_references:
            # References are one-per-line; split on single newlines
            paragraphs = part.split('\n')
        else:
            paragraphs = part.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check if this paragraph has inline math
            if '$' in para_text:
                p = _add_paragraph_with_math(doc, para_text, font_name, font_size,
                                             superscript_citations)
            else:
                p = doc.add_paragraph(style='Normal')
                _add_formatted_runs(p, para_text, font_name, font_size,
                                    superscript_citations=superscript_citations)

            if is_references:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.left_indent = Inches(0.5)
                _suppress_auto_numbering(p)
            elif para_counter == 0 and suppress_first_indent:
                p.paragraph_format.first_line_indent = Pt(0)
            para_counter += 1
            total_words += len(para_text.split())

    return total_words


def generate_docx(input_data):
    journal = load_journal(input_data['journal_profile'])
    author_db = load_authors()

    # Create document — use template.docx by default for consistent formatting.
    # Set "template_path": null in JSON to explicitly skip the template.
    if 'template_path' in input_data and input_data['template_path'] is None:
        template_path = None  # explicitly disabled
    else:
        template_path = input_data.get('template_path')
        if not template_path and os.path.exists(DEFAULT_TEMPLATE):
            template_path = DEFAULT_TEMPLATE
    if template_path:
        doc = create_doc_from_template(template_path)
        # Extract font from template's Normal style (not from journal profile)
        font_name, font_size = _get_template_font(doc)
    else:
        doc = create_doc_from_profile(journal)
        font_name = journal.get('font', 'Times New Roman')
        font_size = journal.get('font_size', 12)

    # -- Title (use Title style if available, else centered bold) --
    # Only emit a title block when one was provided. Otherwise (e.g.,
    # supplementary materials documents that have no title), skip the
    # styled paragraph so the docx doesn't begin with an empty Title.
    if input_data.get('title'):
        try:
            p = doc.add_paragraph(input_data['title'], style='Title')
        except KeyError:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(input_data['title'])
            r.bold = True
            _set_run_font(r, font_name, font_size)

    # -- Authors: ALWAYS use template font (Times New Roman), never journal font --
    # The template Normal style defines the canonical font. Authors, affiliations,
    # and corresponding-author blocks all inherit from it. This ensures consistent
    # typography even when the journal profile specifies a different font (e.g. Arial).
    template_font = _get_template_font(doc)[0] if template_path else font_name

    if input_data.get('authors'):
        authors_numbered, affiliations_numbered = compute_affiliations(
            input_data['authors'], author_db
        )

        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.first_line_indent = Pt(0)
        author_size = font_size  # Same size as body text
        aff_sup_size = 9  # Superscript affiliation numbers

        for i, (name, nums) in enumerate(authors_numbered):
            prefix = ', ' if i > 0 else ''
            r = p.add_run(prefix + name)
            _set_run_font(r, template_font, author_size)
            r = p.add_run(nums)
            _set_run_font(r, template_font, aff_sup_size)
            r.font.superscript = True

        # -- Affiliations (Normal style, no indent) --
        aff_font_size = journal.get('affiliation_font_size', 9)

        for aff in affiliations_numbered:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(aff)
            _set_run_font(r, template_font, aff_font_size)
            r.italic = True

    # -- Corresponding author block (uses template font, same as authors) --
    if input_data.get('corresponding_author'):
        ca = input_data['corresponding_author']
        doc.add_paragraph(style='Normal')  # blank line
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run('*Corresponding author:')
        _set_run_font(r, template_font)
        for line in ca.get('lines', []):
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(line)
            _set_run_font(r, template_font)

    # -- Determine citation rendering --
    cite_style = journal.get('citation_style', '')
    superscript_cites = 'superscript' in cite_style

    # -- Figures directory (for inserting images) --
    figures_dir = input_data.get('figures_dir', '')
    if not figures_dir:
        figures_dir = os.path.dirname(input_data.get('output_path', ''))
    # Root for resolving explicit figure paths from MD ``![](path)``
    manuscript_src_dir = input_data.get('manuscript_src_dir', '') or figures_dir

    # -- Caption/table font sizes --
    caption_font_size = journal.get('caption_font_size', 10)
    table_font_size = journal.get('table_font_size', 9)

    # -- Page break targets (insert page break before these headings) --
    page_break_before = {h.lower() for h in
                         input_data.get('page_break_before', [])}

    # -- Body sections --
    total_words = 0
    for section in input_data['sections']:
        level = section.get('level', 1)
        header = section['header']
        text = section['text']
        sec_type = section.get('type', 'text')

        # Insert page break before specified headings
        if header and header.lower() in page_break_before:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        # Add heading (level 0 means no heading, just text)
        if level > 0 and header:
            add_heading(doc, header, level=level)

        if not text:
            continue

        # --- Table caption ---
        if sec_type == 'table_caption':
            _add_table_caption(doc, text, font_name, font_size=None,
                               superscript_citations=superscript_cites)
            total_words += len(text.split())
            continue

        # --- Table note (below table, Caption style with "Note. " prefix) ---
        if sec_type == 'table_note':
            _add_table_note(doc, text, font_name,
                            font_size=caption_font_size,
                            superscript_citations=superscript_cites)
            total_words += len(text.split())
            continue

        # --- Table data (tab-separated) ---
        if sec_type == 'table':
            _add_table_from_tsv(doc, text, font_name,
                                cell_font_size=table_font_size)
            total_words += len(text.split())
            continue

        # --- Figure + caption (grouped with Square wrapping, or inline) ---
        if sec_type == 'figure':
            fig_num = None
            # Support "FIGURE:N" format (N can be numeric or alphanumeric like S1)
            fig_num_match = re.match(r'^FIGURE:([A-Za-z]?\d+)', text)
            if fig_num_match:
                fig_num = fig_num_match.group(1)
            else:
                # Legacy: [INSERT FIGURE N HERE] format
                fig_match = _INSERT_FIG_RE.search(text)
                if fig_match:
                    fig_num = fig_match.group(1)
            # Extract caption (everything after the FIGURE:N or [INSERT...] line)
            caption_lines = []
            for line in text.split('\n'):
                line = line.strip()
                if (line and not _INSERT_FIG_RE.match(line)
                    and not re.match(r'^FIGURE:[A-Za-z]?\d+', line)):
                    caption_lines.append(line)
            caption_text = ''
            if caption_lines:
                caption_text = ' '.join(caption_lines)
                # Handle **Figure N.** bold markers in caption
                caption_text = re.sub(r'\*\*(Figure [A-Za-z]?\d+\.)\*\*',
                                      r'\1', caption_text)
            fig_width = section.get('figure_width', 6.5)
            # Resolve an explicit path from the MD ``![](path)``
            fig_src = section.get('figure_src')
            fig_path = None
            if fig_src:
                candidate = fig_src if os.path.isabs(fig_src) else os.path.join(
                    manuscript_src_dir, fig_src)
                if os.path.isfile(candidate):
                    fig_path = candidate
            if fig_num and caption_text:
                # Try grouped figure (image + caption as group, Square wrapping)
                try:
                    _insert_grouped_figure(doc, fig_num, caption_text,
                                           figures_dir, font_name, font_size,
                                           superscript_citations=superscript_cites,
                                           target_width_inches=fig_width,
                                           fig_path=fig_path)
                except Exception as e:
                    # Fallback to inline if grouping fails — log the error
                    import traceback
                    print(f"WARNING: Grouped figure failed for Figure {fig_num}, "
                          f"falling back to inline: {e}")
                    traceback.print_exc()
                    _insert_figure(doc, fig_num, figures_dir, font_name, font_size,
                                   fig_path=fig_path)
                    _add_figure_caption(doc, caption_text, font_name,
                                        font_size=font_size,
                                        superscript_citations=superscript_cites)
            elif fig_num:
                _insert_figure(doc, fig_num, figures_dir, font_name, font_size,
                               fig_path=fig_path)
            elif caption_text:
                _add_figure_caption(doc, caption_text, font_name,
                                    font_size=font_size,
                                    superscript_citations=superscript_cites)
            total_words += len(text.split())
            continue

        # --- Table/Figure placeholder (submission version) ---
        if sec_type in ('figure_placeholder', 'table_placeholder'):
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(text)
            _set_run_font(r, font_name, font_size)
            continue

        # --- Figure legends (submission version: at end, Normal style) ---
        if sec_type == 'figure_legends':
            for caption_text in text.split('\n\n'):
                caption_text = caption_text.strip()
                if not caption_text:
                    continue
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.first_line_indent = Pt(0)
                m = re.match(r'(Figure [A-Za-z]?\d+\.)\s*(.*)', caption_text, re.DOTALL)
                if m:
                    r = p.add_run(m.group(1))
                    r.bold = True
                    _set_run_font(r, font_name, font_size)
                    rest = ' ' + m.group(2) if m.group(2) else ''
                    if rest.strip():
                        _add_formatted_runs(p, rest, font_name, font_size,
                                            superscript_citations=superscript_cites)
                else:
                    _add_formatted_runs(p, caption_text, font_name, font_size,
                                        superscript_citations=superscript_cites)
            continue

        # --- Regular body text (with math support) ---
        is_refs = header.lower() in ('references', 'bibliography',
                                     'works cited', 'reference list')
        has_heading = level > 0 and bool(header)

        # Use math-aware paragraph builder if text contains $ delimiters
        if '$' in text:
            wc = add_body_paragraphs_with_math(
                doc, text, font_name, font_size,
                is_references=is_refs,
                superscript_citations=superscript_cites,
                suppress_first_indent=has_heading)
        else:
            wc = add_body_paragraphs(doc, text, font_name, font_size,
                                     is_references=is_refs,
                                     superscript_citations=superscript_cites,
                                     suppress_first_indent=has_heading)
        total_words += wc

    # -- Save --
    output_path = input_data['output_path']
    doc.save(output_path)

    # -- Report --
    word_limit = journal.get('word_limit', 0)
    status = 'OK'
    if word_limit and total_words > word_limit:
        status = f'OVER LIMIT by {total_words - word_limit} words'
    print(f'Saved to {output_path}')
    print(f'Body word count: {total_words}' +
          (f' / {word_limit} limit ({status})' if word_limit else ''))
    for section in input_data['sections']:
        if section.get('text'):
            wc = len(section['text'].split())
            print(f'  {section["header"]}: {wc} words')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        input_data = json.load(sys.stdin)
    else:
        with open(sys.argv[1], encoding='utf-8') as f:
            input_data = json.load(f)
    generate_docx(input_data)
